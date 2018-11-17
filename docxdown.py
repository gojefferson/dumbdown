from docx import Document

from dumbdown import (
    Node,
    TextNode,
    StrongNode,
    ItalNode,
    ParagraphNode,
    DumbDown,
    Tree,
    ReAdapter,
    STRONG_RE,
    ITAL_RE,
)


def extract_first_node(input_string):
    strong = ReAdapter(STRONG_RE, input_string)
    ital = ReAdapter(ITAL_RE, input_string)

    if strong.start_index == 0:
        return DocxStrongNode(content=strong.content), input_string[strong.end_index:]

    if ital.start_index == 0:
        return DocxItalNode(content=ital.content), input_string[ital.end_index:]

    # we have no matches
    if not strong.start_index and not ital.start_index:
        return DocxTextNode(content=input_string), ""

    _max = len(input_string)
    split_point = min(strong.start_index or _max, ital.start_index or _max)
    return DocxTextNode(content=input_string[:split_point]), input_string[split_point:]


class DocxNode(Node):

    def __init__(self, content=""):
        self.content = content
        self.children = []
        if not self.is_leaf_node:
            while content != "":
                node, content = extract_first_node(content)
                self.append_child(node)

    def add_runs(self, p):
        raise NotImplemented("Subclasses of Node must implement this add_run")


class DocxRootNode(DocxNode):

    def write_to_doc(self, doc):
        for child in self.children:
            child.write_to_doc(doc)
        return doc


class DocxTextNode(TextNode):

    def add_runs(self, p):
        p.add_run(self.content)
        return p


class DocxItalNode(ItalNode, DocxNode):

    def add_runs(self, p):
        for child in self.children:
            run = p.add_run(child.content)
            run.italic = True
            if type(child) is DocxStrongNode:
                run.bold = True
        return p


class DocxStrongNode(StrongNode, DocxNode):

    def add_runs(self, p):
        for child in self.children:
            run = p.add_run(child.content)
            run.bold = True
            if type(child) is DocxItalNode:
                run.italic = True
        return p


class DocxParagraphNode(ParagraphNode, DocxNode):

    def write_to_doc(self, doc):
        p = doc.add_paragraph()
        for child in self.children:
            p = child.add_runs(p)


class DocxTree(Tree):

    def __init__(self):
        self.root = DocxRootNode()

    def append_to_doc(self, doc):
        return self.root.write_to_doc(doc)


class DocxDumbDown(DumbDown):

    def __init__(self, md=""):
        self._md = md
        self._tree = DocxTree()
        self._lines = self._md.split("\n")
        self._build_tree()

    def _build_tree(self):
        for line in self._lines:
            p = DocxParagraphNode(content=line)
            self._tree.root.append_child(p)

    def append_to_doc(self, doc):
        return self._tree.append_to_doc(doc)


if __name__ == "__main__":
    document = Document()
    _md = (
        "this is a story about _Hans_\n"
        "*he _IS_ a good dog!*\n"
        "_but_ ... he can also _be *BAD*_\n"
    )
    document = DocxDumbDown(md=_md).append_to_doc(document)
    document.save('demo.docx')
