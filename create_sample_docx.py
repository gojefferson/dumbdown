from docx import Document
from faker import Faker

from dumbdown.docx import confirm_styles_exist_in_document, write_md_to_doc

ESSENTIAL_STYLES = ["Normal", "Heading 1", "Heading 2", "List Paragraph", "No Spacing"]

fake = Faker()


def write_random_paragraphs(document, number=2):
    for i in range(number):
        document.add_paragraph(f"{i}: {fake.text()}")


SAMPLE_MD = (
    "_This project does not use markdown_, but an *extremely simple* formatting language based on Slack. "
    " I call it ... *DUMBDOWN*.\n"
    "It performs in the exact same way that Slack performs with regard to nesting "
    "_*bold* within italics_ and handling various kinds of characters that can be just before or after a "
    "!_*section of formatted text*_:!\n"
    "Because this is _dumb_ -down, we don't do all the that the Slack variant of markdown let's you do, "
    "such as creating quotes "
    "ordered and unordered lists, and code formatting. *Lawyers don't need that stuff*, so we shouldn't "
    "provide it\n"
    "Besides parsing, the project will also render parsed text to _HTML_ and will write it to a *DOCX "
    "document.* This file provides an example of that.\n"
)

if __name__ == "__main__":

    document = Document("legal.docx")

    confirm_styles_exist_in_document(document, ESSENTIAL_STYLES)
    document.add_heading("Sample Legal Document", 0)
    write_md_to_doc(document, SAMPLE_MD, strip_newlines=False, numbered=False)

    document.add_heading("This is a good document", level=1)
    document.add_heading("It has clean formatting", level=2)

    write_md_to_doc(document, SAMPLE_MD, numbered=True)

    write_md_to_doc(document, fake.text(), numbered=True)

    document.add_heading("It uses numbered headings", level=2)
    write_md_to_doc(document, fake.text(), numbered=True)

    document.add_heading("It uses all built-in styles", level=2)
    write_md_to_doc(document, fake.text(), numbered=True)

    document.add_heading("It has numbered paragraphs", level=2)
    write_md_to_doc(document, fake.text(), numbered=True)

    document.add_heading("It needs to be integrated with my markdown thing", level=2)
    write_md_to_doc(document, fake.text(), numbered=True)

    document.add_heading("It can still be improved", level=1)

    document.add_heading("It needs to be integrated with the api repo", level=2)
    write_md_to_doc(document, fake.text(), numbered=True)

    document.save("out/sample.docx")
