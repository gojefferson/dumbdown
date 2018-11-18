from docx import Document
from dumbdown.docxdown import write_md_to_doc


def test_creates_single_paragraph_by_default():
    document = Document()
    write_md_to_doc(document, "first line.\nsecond line\nthird line!")
    assert len(document.paragraphs) == 1


def test_can_create_multiple_paragraphs():
    document = Document()
    write_md_to_doc(document, "first line.\nsecond line\nthird line", strip_newlines=False)
    assert len(document.paragraphs) == 3


def test_correct_runs_correctly():
    # 1. `normal `
    # 2. `ital`
    # 3. ` `
    # 4. `bold`
    document = Document()
    write_md_to_doc(document, "normal _ital_ *bold*")
    assert len(document.paragraphs) == 1
    assert len(document.paragraphs[0].runs) == 4


def test_correct_runs_correctly_nested():
    # 1. `bold `
    # 2. `ital`
    # 3. ` bold`
    document = Document()
    write_md_to_doc(document, "*bold _ital_ bold*")
    assert len(document.paragraphs) == 1
    assert len(document.paragraphs[0].runs) == 3


def test_create_italics_correctly():
    document = Document()
    write_md_to_doc(document, "_ital_")
    assert len(document.paragraphs) == 1
    assert len(document.paragraphs[0].runs) == 1
    run = document.paragraphs[0].runs[0]
    assert run.text == "ital"
    assert run.italic
    assert not run.bold


def test_create_bold_correctly():
    document = Document()
    write_md_to_doc(document, "*bold*")
    assert len(document.paragraphs) == 1
    assert len(document.paragraphs[0].runs) == 1
    run = document.paragraphs[0].runs[0]
    assert run.text == "bold"
    assert not run.italic
    assert run.bold


def test_create_bold_ital_correctly_variant_one():
    document = Document()
    write_md_to_doc(document, "*_bold-ital_*")
    assert len(document.paragraphs) == 1
    assert len(document.paragraphs[0].runs) == 1
    run = document.paragraphs[0].runs[0]
    assert run.text == "bold-ital"
    assert run.italic
    assert run.bold


def test_create_bold_ital_correctly_variant_two():
    document = Document()
    write_md_to_doc(document, "_*bold-ital*_")
    assert len(document.paragraphs) == 1
    assert len(document.paragraphs[0].runs) == 1
    run = document.paragraphs[0].runs[0]
    assert run.text == "bold-ital"
    assert run.italic
    assert run.bold
